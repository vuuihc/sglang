from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "study-notes/推理优化全景-知乎回答粘贴版.md"
OUTPUT = ROOT / "study-notes/推理优化全景-知乎回答粘贴版.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D9E2F0", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color="D9E2F0", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((r_pr, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`|\*([^*]+)\*)"
)


def add_inline(paragraph, text: str, font_size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            if font_size:
                run.font.size = Pt(font_size)
        if match.group(2) is not None:
            add_hyperlink(paragraph, match.group(2), match.group(3))
        elif match.group(4) is not None:
            run = paragraph.add_run(match.group(4))
            run.bold = True
            if font_size:
                run.font.size = Pt(font_size)
        elif match.group(5) is not None:
            run = paragraph.add_run(match.group(5))
            run.font.name = "Sarasa Fixed SC"
            run.font.color.rgb = RGBColor(180, 45, 45)
            run.font.size = Pt((font_size or 10.5) - 0.5)
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Sarasa Fixed SC")
        elif match.group(6) is not None:
            run = paragraph.add_run(match.group(6))
            run.italic = True
            if font_size:
                run.font.size = Pt(font_size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if font_size:
            run.font.size = Pt(font_size)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(35, 41, 47)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 8, "1F5E96"),
        ("Heading 2", 13, 14, 6, "2E74B5"),
        ("Heading 3", 11.5, 10, 4, "315B7D"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(3)


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    total_width = Inches(7.0)
    max_lengths = []
    for col in range(cols):
        values = [len(row[col]) if col < len(row) else 1 for row in rows]
        max_lengths.append(max(4, min(max(values), 28)))
    length_sum = sum(max_lengths)
    widths = [max(0.62, 7.0 * length / length_sum) for length in max_lengths]
    scale = 7.0 / sum(widths)
    widths = [width * scale for width in widths]
    compact_size = 7.2 if cols >= 7 else 8.5 if cols >= 5 else 9.2

    for row_index, row in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(widths[col_index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            text = row[col_index] if col_index < len(row) else ""
            add_inline(paragraph, text, compact_size)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    p_pr.append(shd)
    run = p.add_run(code.rstrip())
    run.font.name = "Sarasa Fixed SC"
    run.font.size = Pt(8.3)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Sarasa Fixed SC")


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            set_paragraph_border(p)
            i += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt, rel_path = image_match.groups()
            image_path = (SOURCE.parent / rel_path).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_together = True
            p.add_run().add_picture(str(image_path), width=Inches(6.7))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(7)
            run = caption.add_run(alt)
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(95, 105, 115)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading_match.group(2))
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "2E74B5")
            p_bdr.append(left)
            p_pr.append(p_bdr)
            add_inline(p, stripped[2:])
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            indent = len(list_match.group(1)) // 2
            style = "List Number" if list_match.group(2)[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.28 + 0.22 * indent)
            add_inline(p, list_match.group(3))
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "> ", "```", "![", "|"))
                or candidate == "---"
                or re.match(r"^(\s*)([-*]|\d+\.)\s+", lines[i])
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("知乎回答复制稿")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(140, 145, 150)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
